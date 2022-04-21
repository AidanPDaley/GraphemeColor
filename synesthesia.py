# Author: Aidan Daley
# Date: 4/20/2022
# Purpose: The following program takes a .docx file as input, and changes all the letters to a corresponding color to mimic grapheme-color synesthesia. The goal of the
# the program is to train associative synesthesia by creating documents that can be read, training the correlation. 


# TODO:
#   - Assign each letter a color. (Different for capitals and lowercase? Probably not)
#   - Read text document
#   - Create a document and start translating
#   - Save document.





from docx import Document
from docx.shared import RGBColor


letterColors = {
    "a": (0xFF, 0x00, 0x00),
	"b": (0x5f, 0x9e, 0xa0),
	"c": (0xff, 0x7e, 0x26),
	"d": (0x07, 0x07, 0x8b),
	"e": (0x5c, 0x5c, 0x5e),
	"f": (0x87, 0x0f, 0x06),
	"g": (0xcf, 0x17, 0xb2),
	"h": (0xf9, 0x81, 0x02),
	"i": (0x60, 0x03, 0x19),
	"j": (0x39, 0x3b, 0xbb),
	"k": (0x3c, 0xee, 0x43),
	"l": (0x0F, 0x55, 0x0F),
	"m": (0xe8, 0x5b, 0x2b),
	"n": (0x0b, 0xa2, 0x06),
	"o": (0x33, 0xb3, 0xa6),
	"p": (0x1c, 0x3c, 0x29),
	"q": (0xf4, 0x91, 0xfc),
	"r": (0x53, 0x00, 0xeb),
	"s": (0xeb, 0x00, 0x2c),
	"t": (0xc4, 0xa3, 0xff),
	"u": (0xb2, 0x04, 0x04),
	"v": (0x70, 0x0b, 0xc3),
	"w": (0xb8, 0x67, 0xb5),
	"x": (0xf4, 0x74, 0x7a),
	"y": (0x4a, 0x4a, 0x4a),
	"z": (0x60, 0x40, 0x03)
}


def graphemeColorLetter(char, docObject):
    charKey = char.lower()
    if charKey in letterColors:
        run = docObject.add_run(char)
        font = run.font
        font.color.rgb = RGBColor(letterColors[charKey][0], letterColors[charKey][1], letterColors[charKey][2])
    else:
        run = docObject.add_run(char)

# Changes color of the letters (prototype)
def graphemeColor(char, docObject):
    if (char == 'a' or char == 'A'):
        run = docObject.add_run(char)
        font = run.font
        font.color.rgb = RGBColor(0xFF,0x00, 0x00)
        #font = run.font.color.rgb(0xFF, 0x00, 0x00)
    else:
        run = docObject.add_run(char)



if __name__ == "__main__":
    sampleText = "Within the first two chapters you are introduced to the green light, the Valley of Ashes, and the billboard from which the eyes of T.J. Eckleburg watch over New York. Within the Great Gatsby, symbolism plays key roles in the characters actions, emotions, and dreams. The Great Gatsby was written by F. Scott Fitzgerald in the early 1920s, while the stock market was at an all time high, directly following the First World War. As the economy thrived, people began to think of the American Dream as something that was easily achievable. This led to corruption and greed. Although a large percentage of the outcome was negative, there were some positive contributions to society, including many technological innovations in manufacturing and every day life. People became shallow and immoral; doing what ever it took to have the newest and best things. They misinterpreted the American Dream to mean that whoever had the best and the most had achieved it and was living far more happily than those who did not have the newest and the most expensive things on the market. F. Scott Fitzgerald was trying to point out the corruption in his book. Fitzgerald uses symbols such as the green light, the Valley of Ashes, and T.J. Eckleburg, to represent the characters corrupt pursuit of the American Dream, materialistic wealth and most of all happiness."
    doc = Document()
    document = doc.add_paragraph()


    for char in sampleText:
        graphemeColorLetter(char, document)


    doc.save('test123.docx')
